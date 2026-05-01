#!/usr/bin/env python3

"""Basic Markdown-to-DOCX exporter for the paper workflow.

Pandoc remains the preferred exporter when available. This fallback handles the
core structures used by the workflow: headings, paragraphs, unordered lists,
ordered lists, and horizontal spacing. It intentionally avoids pretending to be
a full Markdown renderer.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text.strip()


def is_html_comment_line(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("<!--") and stripped.endswith("-->")


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(document: Document, text: str) -> None:
    cleaned = clean_inline(text)
    if cleaned:
        document.add_paragraph(cleaned)


def convert(markdown_path: Path, docx_path: Path) -> None:
    document = Document()
    configure_styles(document)

    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        if buffer:
            add_paragraph(document, " ".join(buffer))
            buffer = []

    for raw_line in markdown_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if is_html_comment_line(line):
            continue

        if not line.strip():
            flush()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush()
            level = min(len(heading.group(1)), 4)
            document.add_heading(clean_inline(heading.group(2)), level=level)
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            flush()
            document.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            flush()
            document.add_paragraph(clean_inline(numbered.group(1)), style="List Number")
            continue

        if line.strip() in {"---", "***", "___"}:
            flush()
            continue

        buffer.append(line.strip())

    flush()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: markdown_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        return 1

    markdown_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])

    if not markdown_path.is_file():
        print(f"Error: markdown file not found: {markdown_path}", file=sys.stderr)
        return 1

    convert(markdown_path, docx_path)
    print(f"Created Word file: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
